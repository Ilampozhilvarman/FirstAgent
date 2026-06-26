import os
import time
import json
import concurrent.futures
import re
from ddgs import DDGS
from openai import OpenAI
from docx import Document
from dotenv import load_dotenv
from playwright.sync_api import sync_playwright

load_dotenv()
client = OpenAI(
    base_url="https://openrouter.ai/api/v1",
    api_key=os.environ.get("OPEN_ROUTER_API_KEY")
)

def _call_llm(messages, response_format=None):
    models = ["openrouter/free"]
    for model in models:
        for attempt in range(3):
            try:
                kwargs = {"model": model, "messages": messages}
                if response_format:
                    kwargs["response_format"] = response_format
                response = client.chat.completions.create(**kwargs)
                return response.choices[0].message.content
            except Exception:
                time.sleep(2 * (attempt + 1))
    raise Exception("All endpoints failed after retries")

def _scrape_url_with_playwright(url):
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            context = browser.new_context(user_agent="Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36")
            page = context.new_page()
            page.goto(url, timeout=10000, wait_until="networkidle")
            text_content = page.locator("p, h1, h2, h3, li").all_inner_texts()
            browser.close()
            full_text = " ".join([t.strip() for t in text_content if t.strip()])
            return full_text[:5000] 
    except Exception:
        return ""

def _search_and_extract(query, max_results=3):
    search_results = []
    try:
        with DDGS() as ddgs:
            for r in ddgs.text(query, max_results=max_results):
                if not isinstance(r, dict) or "href" not in r:
                    continue
                url = r["href"]
                title = r.get("title", "Web Source")
                page_text = _scrape_url_with_playwright(url)
                if page_text:
                    search_results.append({
                        "title": title,
                        "url": url,
                        "content": page_text
                    })
    except Exception:
        pass
    return search_results

def _fetch_all_searches_concurrently(queries):
    all_results = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
        futures = {executor.submit(_search_and_extract, q): q for q in queries}
        for future in concurrent.futures.as_completed(futures):
            all_results.extend(future.result())
    return all_results

def _format_results(results):
    seen_urls = set()
    text = ""
    counter = 1
    for r in results:
        if r['url'] in seen_urls:
            continue
        seen_urls.add(r['url'])
        text += f"\n[{counter}] {r['title']}\nURL: {r['url']}\nFull Content: {r['content']}\n"
        counter += 1
    return text

def _generate_dynamic_queries(task):
    prompt = f"""You are a research planning assistant. Analyze the user request and generate exactly 3 distinct, targeted search engine queries that will gather the most comprehensive, balanced data.
    
USER REQUEST: {task}

Return ONLY a valid JSON array of strings. Do not wrap it in markdown code blocks.
Example Output format: ["query 1", "query 2", "query 3"]"""

    messages = [{"role": "user", "content": prompt}]
    try:
        raw_json = _call_llm(messages)
        clean_json = raw_json.strip().lstrip("```json").rstrip("```").strip()
        queries = json.loads(clean_json)
        if isinstance(queries, list) and len(queries) > 0:
            return queries
    except Exception:
        pass
    return [task, f"{task} deep dive", f"{task} analysis"]

def _write_initial_draft(task, context):
    messages = [{
        "role": "user",
        "content": f"You are a research agent.\n\nDEEP WEB CONTENT LAYER:\n{context}\n\nTASK:\n{task}\n\nWrite a structured initial draft of a report covering:\n- Overview\n- Key Findings\n- Pros and Cons\n- Final Recommendation"
    }]
    return _call_llm(messages)

def _refine_report(task, context, initial_draft):
    messages = [{
        "role": "user",
        "content": f"You are an expert editor and research analyst.\n\nORIGINAL USER TASK:\n{task}\n\nDEEP WEB CONTENT LAYER:\n{context}\n\nINITIAL DRAFT GENERATED:\n{initial_draft}\n\nCRITIQUE AND REFINE TASK:\nReview the initial draft. Check it against the comprehensive text to ensure structural accuracy, eliminate fluff, and fix gaps.\nProvide a finalized, comprehensive, and highly polished version of the report using crisp Markdown formatting. Keep the clean section headers:\n- Overview\n- Key Findings\n- Pros and Cons\n- Final Recommendation"
    }]
    return _call_llm(messages)

def run_research_pipeline(task):
    queries = _generate_dynamic_queries(task)
    raw_results = _fetch_all_searches_concurrently(queries)
    context = _format_results(raw_results)
    draft = _write_initial_draft(task, context)
    final_report = _refine_report(task, context, draft)
    return final_report

def save_report_to_docx(task, report_content, filename):
    doc = Document()
    doc.add_heading(task, 0)
    lines = report_content.split('\n')
    for line in lines:
        cleaned = line.strip()
        if not cleaned:
            continue
        if cleaned.startswith('# '):
            doc.add_heading(cleaned.lstrip('# ').strip(), level=1)
        elif cleaned.startswith('## '):
            doc.add_heading(cleaned.lstrip('## ').strip(), level=2)
        elif cleaned.startswith('### '):
            doc.add_heading(cleaned.lstrip('### ').strip(), level=3)
        elif cleaned.startswith('- ') or cleaned.startswith('* '):
            text_body = cleaned[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text_body)
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, cleaned)
    doc.save(filename)

def _add_formatted_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)