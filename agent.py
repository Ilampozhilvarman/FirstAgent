import os
import re
import subprocess
from datetime import datetime
import time
from ddgs import DDGS
from openai import OpenAI
from docx import Document
from dotenv import load_dotenv

load_dotenv()
client = OpenAI(
    base_url="https://openrouter.ai/api/v1",
    api_key=os.environ.get("OPEN_ROUTER_API_KEY")
)

os.makedirs("reports", exist_ok=True)

def search_web(query, max_results=5):
    results = []
    with DDGS() as ddgs:
        for r in ddgs.text(query, max_results=max_results):
            if not isinstance(r, dict):
                continue
            if "title" not in r or "href" not in r:
                continue
            results.append({
                "title": r["title"],
                "url": r["href"],
                "snippet": r.get("body", "")
            })
    return results

def call_llm(messages):
    models = ["openrouter/free"]
    for model in models:
        for attempt in range(3):
            try:
                response = client.chat.completions.create(
                    model=model,
                    messages=messages
                )
                return response.choices[0].message.content
            except Exception as e:
                print("Model failed:", model)
                print("Attempt:", attempt + 1)
                print(e)
                time.sleep(5 * (attempt + 1))
    raise Exception("All models failed after retries")

def format_results(results):
    text = ""
    for i, r in enumerate(results, 1):
        text += f"""
[{i}] {r['title']}
URL: {r['url']}
Snippet: {r['snippet']}
"""
    return text

def write_report(task, context):
    messages = [{
        "role": "user",
        "content": f"""
You are a research agent.

WEB RESULTS:
{context}

TASK:
{task}

Write a structured report:
- Overview
- Key Findings
- Pros and Cons
- Final Recommendation
"""
    }]

    return call_llm(messages)
while True:
    task = input("\nTask: ")
    if task.lower() == "quit":
        print("Quitting...")
        print("Come again later!")
        break
    print("Searching web...")
    queries = [
        task,
        task + " reviews",
        task + " comparison",
        task + " pros and cons"
    ]
    all_results = []
    for q in queries:
        all_results.extend(search_web(q))
    context = format_results(all_results)
    print("Writing report...")
    report = write_report(task, context)
    safe_name = re.sub(r'[^a-zA-Z0-9_ ]', '', task)
    safe_name = safe_name.strip().replace(" ", "_")[:40]
    filename = f"reports/{safe_name}.docx"
    doc = Document()
    doc.add_heading(task, 0)
    doc.add_paragraph(report)
    doc.save(filename)
    print(f"\nSaved: {filename}")
    subprocess.run(["open", filename])